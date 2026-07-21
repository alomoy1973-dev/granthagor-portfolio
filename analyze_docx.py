import docx
import json
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

def extract_poem_titles_from_docx(filename):
    """Extract poem titles - look for lines that are bold/heading style or match numbered patterns"""
    doc = docx.Document(filename)
    poems = []
    current_title = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is a heading/title - bold text or heading style
        is_title = False
        if para.style.name.startswith('Heading'):
            is_title = True
        elif para.runs and all(run.bold for run in para.runs if run.text.strip()):
            is_title = True
        
        if is_title and text:
            if current_title:
                poems.append({'title': current_title, 'content': current_content[:]})
            current_title = text
            current_content = []
        else:
            current_content.append(text)
    
    if current_title:
        poems.append({'title': current_title, 'content': current_content[:]})
    
    return poems

def extract_all_paragraphs(filename):
    doc = docx.Document(filename)
    results = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        if text:
            results.append({'text': text, 'style': style, 'bold': is_bold})
    return results

print("=== TINNOMURI - Paragraph Styles ===")
tinn_paras = extract_all_paragraphs("তিন্নোমুরি (TINNOMURI).docx")
print(f"Total non-empty paragraphs: {len(tinn_paras)}")
# Show first 30 to understand structure
for p in tinn_paras[:30]:
    marker = "[BOLD]" if p['bold'] else "      "
    print(f"{marker} [{p['style']}] {p['text'][:60]}")

print()
print("=== MONPUDI - Paragraph Styles ===")
mono_paras = extract_all_paragraphs("মনপুদি.docx")
print(f"Total non-empty paragraphs: {len(mono_paras)}")
for p in mono_paras[:30]:
    marker = "[BOLD]" if p['bold'] else "      "
    print(f"{marker} [{p['style']}] {p['text'][:60]}")
