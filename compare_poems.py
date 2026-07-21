import docx
import json
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

# =============================================
# Load website poems
# =============================================
with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)
website_poems = [x for x in data if x.get('category') == 'poem']
website_titles = [p['title'].strip() for p in website_poems]
print(f"Website has {len(website_poems)} poems")
print()

# =============================================
# MONPUDI - Extract titles from TOC (toc 1 style)
# =============================================
mono_doc = docx.Document("মনপুদি.docx")
monpudi_titles = []
for para in mono_doc.paragraphs:
    if para.style.name.startswith('toc'):
        text = para.text.strip()
        if '\t' in text:
            title = text.split('\t')[0].strip()
        else:
            title = re.sub(r'[\.\s\d]+$', '', text).strip()
        if title and title not in ['মনপুদি', 'আলোময় চাকমা']:
            monpudi_titles.append(title)

print(f"=== মনপুদি: {len(monpudi_titles)} poems ===")
for i, t in enumerate(monpudi_titles, 1):
    # Check if on website
    found = any(t in wt or wt in t for wt in website_titles)
    mark = "✓" if found else "❌ MISSING"
    print(f"  {i:2d}. {mark} | {t}")

print()

# =============================================
# TINNOMURI - Extract titles
# Titles in this docx are standalone lines that are NOT part of poem body
# They appear before multi-line stanzas
# Strategy: a title is a short line that doesn't end in comma/period/special char
# and the NEXT paragraph is a continuation of a poem line
# =============================================
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
all_tinn = []
for para in tinn_doc.paragraphs:
    text = para.text.strip()
    if text:
        all_tinn.append(text)

# Find titles: lines that look like titles (short, no punctuation at end, 
# followed by poem-body lines, and dates like DD/MM/YYYY mark end of poem)
tinn_titles = []
skip_first = True  # skip book title
date_pattern = re.compile(r'\d{2}/\d{2}/\d{4}')

i = 0
while i < len(all_tinn):
    line = all_tinn[i]
    
    # Skip the book title and author name
    if skip_first and line in ['তিন্নোমুরি', 'আলোময় চাকমা', 'তিন্নোমুরি (TINNOMURI)']:
        i += 1
        continue
    skip_first = False
    
    # A title is likely a standalone line that:
    # 1. Does NOT end with typical poem punctuation (। , ; ?)
    # 2. Is not a date
    # 3. Is relatively short
    # 4. Is NOT a multi-line continuation (contains newlines = poem line)
    is_date = bool(date_pattern.search(line))
    ends_with_punct = line.endswith(('।', ',', ';', '!', '?', '।।', '।।।'))
    is_multiline = '\n' in line
    is_short = len(line) < 40
    
    # Heuristic: if this is a non-punctuated short line followed by a comma/period-ending line
    if (not is_date and not ends_with_punct and not is_multiline and is_short 
            and i + 1 < len(all_tinn)):
        next_line = all_tinn[i+1] if i+1 < len(all_tinn) else ''
        # Next line should look like a poem body (ends with comma, period etc)
        next_looks_like_poem = (next_line.endswith((',', '।', ';', '!', '?')) 
                                or '\n' in next_line
                                or len(next_line) > 30)
        if next_looks_like_poem:
            tinn_titles.append(line)
    i += 1

print(f"=== তিন্নোমুরি: {len(tinn_titles)} poems detected ===")
for i, t in enumerate(tinn_titles, 1):
    found = any(t in wt or wt in t for wt in website_titles)
    mark = "✓" if found else "❌ MISSING"
    print(f"  {i:2d}. {mark} | {t}")
