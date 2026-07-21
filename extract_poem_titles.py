import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

# =============================================
# MONPUDI - Extract titles from TOC
# =============================================
print("=== MONPUDI (মনপুদি) - Poem Titles from TOC ===")
mono_doc = docx.Document("মনপুদি.docx")
monpudi_titles = []
for para in mono_doc.paragraphs:
    if para.style.name.startswith('toc'):
        text = para.text.strip()
        # Remove page numbers
        if '\t' in text:
            title = text.split('\t')[0].strip()
        else:
            # Remove trailing dots and numbers
            title = text.rstrip('0123456789. \t').strip()
        if title and title != 'মনপুদি':
            monpudi_titles.append(title)
            print(f"  {len(monpudi_titles)}. {title}")

print(f"\nTotal poems in মনপুদি: {len(monpudi_titles)}")

# =============================================
# TINNOMURI - Extract titles (all Normal style, need to detect by content pattern)
# =============================================
print("\n=== TINNOMURI (তিন্নোমুরি) - Poem Titles ===")
print("(Detecting titles - these appear as standalone lines before poem content)")
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
tinn_paras = []
for para in tinn_doc.paragraphs:
    text = para.text.strip()
    if text:
        tinn_paras.append(text)

# Print all paragraphs to manually identify pattern
print("ALL PARAGRAPHS (first 200):")
for i, t in enumerate(tinn_paras[:200]):
    print(f"  [{i:3d}] {t[:80]}")
