import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)

existing_titles = set(x['title'].strip() for x in data if x.get('category') == 'poem')
print(f"Current poems: {len(existing_titles)}")

def process_lines(raw_lines):
    result = []
    for raw in raw_lines:
        for line in raw.split('\n'):
            line = line.strip()
            if not line:
                if result and result[-1] != '__STANZA__':
                    result.append('__STANZA__')
            else:
                # Skip date-only lines at end
                if len(result) > 0:
                    result.append(line)
                else:
                    result.append(line)
    while result and result[-1] == '__STANZA__':
        result.pop()
    return result

def add_poem(poem_id, title, content_lines, source):
    content = process_lines(content_lines)
    excerpt = content[0] if content else title
    new_poem = {
        "id": poem_id,
        "title": title,
        "category": "poem",
        "badge": "কবিতা",
        "excerpt": excerpt[:100],
        "date": "তারিখ অজানা",
        "readTime": f"{max(1, len(content)//15)} মিনিট পাঠ",
        "isFeatured": False,
        "content": content,
        "source": source
    }
    data.append(new_poem)
    print(f"  ✅ {title} ({len(content)} lines)")
    return new_poem

# =============================================
# TINNOMURI - Using exact paragraph ranges
# =============================================
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
tp = [p.text for p in tinn_doc.paragraphs]

# Map: (title, start_para, end_para, poem_id)
tinn_sections = [
    ("ওলি গীত ( ঘুমপাড়ানি গান)", 86, 95, "poem-tn-n1"),
    ("বিঝু বেড়ান্", 116, 128, "poem-tn-n2"),
    ("কামানি রুয়ে", 188, 194, "poem-tn-n3"),
    ("আয় তুঙ্গোবী", 194, 205, "poem-tn-n4"),
    ("কলিয়ে ভরমেল' পিত্থিমী", 261, 268, "poem-tn-n5"),
    ("অফিস পাড়া", 268, 276, "poem-tn-n6"),
    ("নাদঙ্ ছাড়া", 419, 429, "poem-tn-n7"),
    ("উং চুয়া প্যুঅঃ", 520, 525, "poem-tn-n8"),
]

print("\nAdding from তিন্নোমুরি:")
for title, start, end, pid in tinn_sections:
    if title not in existing_titles:
        # Content is from start+1 to end (exclusive)
        content_lines = tp[start+1:end]
        add_poem(pid, title, content_lines, "তিন্নোমুরি")
    else:
        print(f"  ⏭ Skip (already exists): {title}")

# =============================================
# MONPUDI - হোচপাঙ্ ভিলিনে
# Content is embedded in the same paragraph as title (with \n)
# =============================================
mono_doc = docx.Document("মনপুদি.docx")
mp = [p for p in mono_doc.paragraphs if not p.style.name.startswith('toc')]
mp_text = [p.text for p in mp]

print("\nAdding from মনপুদি:")
# Find হোচপাঙ্ ভিলিনে
for i, text in enumerate(mp_text):
    if 'হোচপাঙ্ ভিলিনে' in text:
        # Get content from this para onward until next title
        content_from_same_para = text.replace('হোচপাঙ্ ভিলিনে', '').strip()
        content_lines = []
        if content_from_same_para:
            content_lines.append(content_from_same_para)
        # Collect next paras until a known title
        mono_known_titles = [
            "তর্ মর্ বানাহ্ দ্বিজনর্", "তরে স্ববনে দেগিম্ ভিলি",
            "হিলোভিদিরে", "হোচপানার গঙাঁরে"
        ]
        j = i + 1
        while j < len(mp_text):
            if any(mt in mp_text[j] for mt in mono_known_titles):
                break
            content_lines.append(mp_text[j])
            j += 1
        
        title = "হোচপাঙ্ ভিলিনে"
        if title not in existing_titles:
            add_poem("poem-mp-n1", title, content_lines, "মনপুদি")
        break

# Save
with open('writings.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, separators=(',', ':'))

poems_total = [x for x in data if x.get('category') == 'poem']
print(f"\n✅ Total poems now: {len(poems_total)}")
print(f"Total entries: {len(data)}")
