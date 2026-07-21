import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)

# Remove the bad হোচপাঙ্ entry (2143 lines) and keep the search going
data = [x for x in data if not (x.get('title') == 'হোচপাঙ্ ভিলিনে' and len(x.get('content', [])) > 100)]
print(f"After cleanup: {len(data)} entries")

def process_lines(raw_lines):
    """Process raw text lines into clean content with stanza markers"""
    result = []
    for raw in raw_lines:
        for line in raw.split('\n'):
            line = line.strip()
            if not line:
                if result and result[-1] != '__STANZA__':
                    result.append('__STANZA__')
            else:
                result.append(line)
    while result and result[-1] == '__STANZA__':
        result.pop()
    return result

def extract_by_paragraph_index(filename, title_para_idx_map):
    """Extract poems given the paragraph indices of their titles"""
    doc = docx.Document(filename)
    paras = list(doc.paragraphs)
    
    # Sort by index
    sorted_titles = sorted(title_para_idx_map.items(), key=lambda x: x[1])
    all_indices = [idx for _, idx in sorted_titles]
    
    poems = {}
    for i, (title, start_idx) in enumerate(sorted_titles):
        # Content is from start_idx+1 to the next title index
        end_idx = all_indices[i+1] if i+1 < len(all_indices) else len(paras)
        content_paras = [paras[j].text for j in range(start_idx+1, end_idx)]
        poems[title] = content_paras
    
    return poems

# =============================================
# TINNOMURI - precise paragraph indices from debug output
# =============================================
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
tinn_paras = list(tinn_doc.paragraphs)

# Find exact paragraph indices by scanning
tinn_title_indices = {}
tinn_targets = [
    "ওলি গীত ( ঘুমপাড়ানি গান)",
    "বিঝু বেড়ান্",
    "কামানি রুয়ে",
    "আয় তুঙ্গোবী",
    "কলিয়ে ভরমেল' পিত্থিমী",
    "অফিস পাড়া",
    "নাদঙ্ ছাড়া",
    "উং চুয়া প্যুঅঃ",
]

for i, para in enumerate(tinn_paras):
    t = para.text.strip()
    for title in tinn_targets:
        if t == title or t == title.strip():
            tinn_title_indices[title] = i
            print(f"Found '{title}' at para {i}")

# For বিঝু বেড়ান্ - find in full text
for i, para in enumerate(tinn_paras):
    t = para.text.strip()
    if 'বিঝু বেড়' in t:
        print(f"  'বিঝু বেড়' found at {i}: {repr(t[:50])}")
    if 'আয় তুঙ' in t:
        print(f"  'আয় তুঙ' found at {i}: {repr(t[:50])}")
    if 'কলিয়ে' in t and len(t) < 30:
        print(f"  'কলিয়ে' found at {i}: {repr(t[:50])}")
    if 'সাবালা' in t and len(t) < 30:
        print(f"  'সাবালা' at {i}: {repr(t[:50])}")
    if 'ওলি গীত' in t and len(t) < 40:
        print(f"  'ওলি গীত' at {i}: {repr(t[:50])}")
