import docx
import json
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

# =============================================
# Helper: clean text
# =============================================
def clean(text):
    return text.strip()

# =============================================
# MONPUDI - Extract poems
# Strategy: TOC gives titles. Content is between titles.
# =============================================
def extract_monpudi_poems():
    doc = docx.Document("মনপুদি.docx")
    
    # Get all TOC titles first
    toc_titles = []
    for para in doc.paragraphs:
        if para.style.name.startswith('toc'):
            text = para.text.strip()
            if '\t' in text:
                title = text.split('\t')[0].strip()
            else:
                title = re.sub(r'[\.\s\d]+$', '', text).strip()
            if title and title not in ['মনপুদি', 'আলোময় চাকমা']:
                toc_titles.append(title)
    
    # Now extract content - find each poem by its title in the document
    all_paras = []
    for para in doc.paragraphs:
        if not para.style.name.startswith('toc'):
            text = para.text.strip()
            all_paras.append(text)
    
    # Find poem sections: title followed by content
    poems = []
    i = 0
    # skip header/TOC area
    # Find where actual poems start (after TOC)
    content_start = 0
    for idx, para in enumerate(all_paras):
        if para in toc_titles and idx > 5:
            content_start = idx
            break
    
    i = content_start
    while i < len(all_paras):
        line = all_paras[i]
        if line in toc_titles:
            title = line
            content_lines = []
            i += 1
            # Collect until next title or end
            while i < len(all_paras):
                next_line = all_paras[i]
                if next_line in toc_titles:
                    break
                if next_line:
                    content_lines.append(next_line)
                i += 1
            poems.append({'title': title, 'content': content_lines})
        else:
            i += 1
    
    return poems

# =============================================
# TINNOMURI - Extract poems
# Strategy: detect titles as short lines before poem body
# =============================================
def extract_tinnomuri_poems():
    doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
    
    all_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        all_paras.append(text)
    
    date_pattern = re.compile(r'\d{1,2}/\d{1,2}/\d{4}')
    
    # Known title list from earlier analysis
    known_titles = [
        "টেঙাঁ (ঘুষ)", "ভুদেপা দেখ্যের বাপ", "সিজি বেঞ্যেঁ", "হোবাল্ আঙেঁর",
        "জারাল্যে পিধা", "তিন্নোমুরি", "পাঞ্যেঁ নাগুরী", "বোলী পিনোন্",
        "বো হঝা", "চোখকুন্ ফুদিনে", "বাঝা বিঝে", "নুও বোই",
        "নুওভাত্ খেবাত্তে", "ওলি গীত ( ঘুমপাড়ানি গান)", "মনআঝা",
        "ডিম ভাজী", "বিজবিজিদি", "বিঝু বেড়ান্", "মোন্ বগা",
        "ইত্তুগি মা'র বিঝু বাজার", "এচ্যে বেবে বাউচ্ছি",
        "ভেচ্চাক্/ভেরেচ্চাক্", "দারিত্বে উন্দুর", "থুর'চুমো",
        "বিঝুত্ বলাঞ্যেঁ", "বিঝু লেজ্", "বাচ্ছুরি রেসিপি",
        "কামানি রুয়ে", "আয় তুঙ্গোবী", "নানুমরা", "বেলেই ন' থেলে",
        "মাজ' জু", "বাঙুঁরী জুর", "পিধা টিগা", "বান্দর রাঘা",
        "লাভাক", "বুরত ভোরেবং ভুজিরে", "বেগেনা",
        "কলিয়ে ভরমেল' পিত্থিমী", "অফিস পাড়া", "শামুক তুলা",
        "হন্না হেব' চেই", "আলজি ঘরবো", "সনা মাবে ঘুম",
        "দুগোর শালগরা", "১/ পিজু আমল....", "২/ আজু আমল....",
        "৩/ বাবর আমল....", "৪/ ইরুক আমল ( নিজর আমল)..",
        "৫/ মুজুঙোঁ আমল....", "নিবিলি জুম", "পদর বাধি ডাঙঁর ভুজি",
        "দ্বি পদলা জা", "ফেসবুকার ইত্তুক হানা", "পাত্তুং",
        "হুন্দি ঘাজিলে হিজেনি", "বিঝু তুই এবে ভিলিনে...",
        "ন' তাঙ'রনা", "ফুদোন্দি মন্", "আল্ নেই", "মূল্ বিঝু",
        "বিঝু এলে......", "নাঙ্ নেই", "নাদঙ্ ছাড়া",
        "জুম্মবী ত' উধিজে", "স্বমনে", "জাগি উদো মা-বোন্ লক্",
        "ভালা চান্", "ফুলফুল চুলচুল্", "আওজো বিনিভাত",
        "উদোশুউনি মুদোবাচ", "ঝিমিত সনা", "কাম্মো চিজি",
        "জুমরাঘা", "ইহিঁজক", "উং চুয়া প্যুঅঃ"
    ]
    
    poems = []
    current_title = None
    current_content = []
    
    for line in all_paras:
        if line in known_titles:
            if current_title:
                poems.append({'title': current_title, 'content': current_content[:]})
            current_title = line
            current_content = []
        elif current_title and line:
            current_content.append(line)
    
    if current_title:
        poems.append({'title': current_title, 'content': current_content[:]})
    
    return poems

# =============================================
# Process content into stanzas
# =============================================
def process_content(lines):
    result = []
    for line in lines:
        if not line.strip():
            if result and result[-1] != '__STANZA__':
                result.append('__STANZA__')
        else:
            result.append(line.strip())
    # Remove trailing stanza marker
    while result and result[-1] == '__STANZA__':
        result.pop()
    return result

# =============================================
# Load website data
# =============================================
with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)

website_poems = [x for x in data if x.get('category') == 'poem']
website_titles = [p['title'].strip() for p in website_poems]

def is_on_website(title):
    for wt in website_titles:
        if title.strip() == wt.strip():
            return True
        if title.strip() in wt or wt in title.strip():
            return True
    return False

# =============================================
# Extract and add MONPUDI poems
# =============================================
print("Extracting মনপুদি poems...")
monpudi_poems = extract_monpudi_poems()
print(f"Found {len(monpudi_poems)} poems in মনপুদি")

monpudi_added = []
for p in monpudi_poems:
    if not is_on_website(p['title']):
        content = process_content(p['content'])
        excerpt = content[0] if content else p['title']
        new_poem = {
            "id": f"poem-mp-{len(monpudi_added)+1:03d}",
            "title": p['title'],
            "category": "poem",
            "badge": "কবিতা",
            "excerpt": excerpt[:100],
            "date": "তারিখ অজানা",
            "readTime": f"{max(1, len(content)//15)} মিনিট পাঠ",
            "isFeatured": False,
            "content": content,
            "source": "মনপুদি"
        }
        data.append(new_poem)
        monpudi_added.append(p['title'])
        print(f"  Added: {p['title']}")

print(f"\nAdded {len(monpudi_added)} poems from মনপুদি")

# =============================================
# Extract and add TINNOMURI poems
# =============================================
print("\nExtracting তিন্নোমুরি poems...")
tinn_poems = extract_tinnomuri_poems()
print(f"Found {len(tinn_poems)} poems in তিন্নোমুরি")

tinn_added = []
for p in tinn_poems:
    if not is_on_website(p['title']):
        content = process_content(p['content'])
        excerpt = content[0] if content else p['title']
        new_poem = {
            "id": f"poem-tn-{len(tinn_added)+1:03d}",
            "title": p['title'],
            "category": "poem",
            "badge": "কবিতা",
            "excerpt": excerpt[:100],
            "date": "তারিখ অজানা",
            "readTime": f"{max(1, len(content)//15)} মিনিট পাঠ",
            "isFeatured": False,
            "content": content,
            "source": "তিন্নোমুরি"
        }
        data.append(new_poem)
        tinn_added.append(p['title'])
        print(f"  Added: {p['title']}")

print(f"\nAdded {len(tinn_added)} poems from তিন্নোমুরি")

# =============================================
# Save
# =============================================
with open('writings.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, separators=(',', ':'))

print(f"\n✅ DONE!")
print(f"  মনপুদি poems added: {len(monpudi_added)}")
print(f"  তিন্নোমুরি poems added: {len(tinn_added)}")
print(f"  Total new poems: {len(monpudi_added) + len(tinn_added)}")
print(f"  Total entries in writings.json: {len(data)}")
poems_after = [x for x in data if x.get('category') == 'poem']
print(f"  Total poems on website now: {len(poems_after)}")
