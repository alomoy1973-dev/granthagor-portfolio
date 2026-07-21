import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)

def process_content(lines):
    result = []
    for line in lines:
        if not line.strip():
            if result and result[-1] != '__STANZA__':
                result.append('__STANZA__')
        else:
            result.append(line.strip())
    while result and result[-1] == '__STANZA__':
        result.pop()
    return result

# =============================================
# Manually add the 9 remaining poems
# by reading from the DOCX directly
# =============================================

# Read all paragraphs from both docs
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
mono_doc = docx.Document("মনপুদি.docx")

tinn_all = [p.text for p in tinn_doc.paragraphs]
mono_all = [p.text for p in mono_doc.paragraphs if not p.style.name.startswith('toc')]

def extract_section(all_paras, start_title, end_titles):
    """Extract content between start_title and the next known title"""
    collecting = False
    content = []
    for line in all_paras:
        text = line.strip()
        if text == start_title:
            collecting = True
            continue
        if collecting:
            if text in end_titles:
                break
            if text:
                content.append(text)
    return content

# All known title markers in tinnomuri for boundary detection
tinn_titles_all = [
    "টেঙাঁ (ঘুষ)", "ভুদেপা দেখ্যের বাপ", "সিজি বেঞ্যেঁ", "হোবাল্ আঙেঁর",
    "জারাল্যে পিধা", "তিন্নোমুরি", "পাঞ্যেঁ নাগুরী", "বোলী পিনোন্",
    "বো হঝা", "চোখকুন্ ফুদিনে", "বাঝা বিঝে", "নুও বোই",
    "নুওভাত্ খেবাত্তে", "ওলি গীত ( ঘুমপাড়ানি গান)", "মনআঝা",
    "ডিম ভাজী", "বিজবিজিদি", "বিঝু বেড়ান্", "মোন্ বগা",
    "ইত্তুগি মা'র বিঝু বাজার", "এচ্যে বেবে বাউচ্ছি",
    "ভেচ্চাক্/ভেরেচ্চাক্", "দারিত্বে উন্দুর", "থুর'চুমো",
    "বিঝুত্ বলাঞ্যেঁ", "বিঝু লেজ্", "বাচ্ছুরি রেসিপি",
    "কামানি রুয়ে", "আয় তুঙ্গোবী", "নানুমরা", "বেলেই ন' থেলে",
    "মাজ' জু", "বাঙুঁরী জুর", "পিধা টিগা", "বান্দর রাঘা",
    "লাভাক", "বুরত ভোরেবং ভুজিরে", "সাবালা কন্না অব'?", "বেগেনা",
    "কলিয়ে ভরমেল' পিত্থিমী", "অফিস পাড়া", "শামুক তুলা",
    "হন্না হেব' চেই", "আলজি ঘরবো", "সনা মাবে ঘুম",
    "দুগোর শালগরা", "১/ পিজু আমল....", "২/ আজু আমল....",
    "৩/ বাবর আমল....", "৪/ ইরুক আমল ( নিজর আমল)..",
    "৫/ মুজুঙোঁ আমল....", "নিবিলি জুম", "পদর বাধি ডাঙঁর ভুজি",
    "দ্বি পদলা জা", "ফেসবুকার ইত্তুক হানা", "পাত্তুং",
    "হুন্দি ঘাজিলে হিজেনি", "বিঝু তুই এবে ভিলিনে...",
    "ন' তাঙ'রনা", "ফুদোন্দি মন্", "আল্ নেই", "মূল্ বিঝু",
    "বিঝু এলে......", "নাঙ্ নেই", "নাদঙ্ ছাড়া",
    "জুম্মবী ত' উধিজে", "স্বমনে", "জাগি উদো মা-বোন্ লক্",
    "ভালা চান্", "ফুলফুল চুলচুল্", "আওজো বিনিভাত",
    "উদোশুউনি মুদোবাচ", "ঝিমিত সনা", "কাম্মো চিজি",
    "জুমরাঘা", "ইহিঁজক", "উং চুয়া প্যুঅঃ"
]

mono_titles_all = [
    "মনপুদি", "হোচপাঙ্ ভিলিনে", "তর্ মর্ বানাহ্ দ্বিজনর্",
    "তরে স্ববনে দেগিম্ ভিলি", "হিলোভিদিরে", "হোচপানার গঙাঁরে",
    "জুম্মবী পরানি মর", "সংপোর হোচপানা", "দিল' হেনা",
    "তর পূন্নিমা মর আঙুঁচ্যা", "এক আজুল হোচপানা",
    "মনপুদির হোচপানা", "ও মর নুদিবানী", "মিধে আঝি",
    "বাচ্ছে থাক্", "ভাবনার রাধামন-ধনপুদি", "বারিজার এ দিনোত্",
    "ঈধোত্ উধে", "কুজ্", "সমার (বন্ধু) মরে হোচপেচ্",
    "হোচপানার চারাগাছ", "চিত্তিবী", "মুই আগঙ ত' লগে",
    "জাগুলুক্ মন্.....", "ঈধোত রাগেচ", "বুঝি লোচ পরানি",
    "কাট্টোন", "ভাবনাত ভাজদে গম লাগে", "হারে হোই হোচপানা",
    "আড়ন্দি রাজা দেজত জুন্মিলে", "মুই আগঙ ত' লগে",
    "বানি থোই দিস", "হোচপানার তিরাজ", "জিংহানির লামা",
    "এক গোদেল হোচপানা", "যুনি ফেলে ন' যাস পরানি",
    "তরে বলাঙঁর", "রিত মংঘাহ হোচপানা", "তুই এবে ভিলি",
    "একফুদো পহন হোচপানা", "হোচপানা কারে কয়",
    "হোচপানার লামা", "মেইয়্যাঁ লামা", "অরাযুগ হোচপানা",
    "ঈদোঘর", "মরা স্বমন", "চিগোন আওজ",
    "যুনি পাং হোচপানা তর", "কমলে নিবেগি?",
    "উয়ুরী উদিবার ভারী ধাব", "ফোন", "চ্যাটিং",
    "স্বমনানি পুড়িযার", "আগ সজপদর", "স্বমন"
]

# Add remaining 9 poems
remaining_poems = [
    # (title, source_doc_paras, all_titles_for_boundary, source_book)
    ("হোচপাঙ্ ভিলিনে", mono_all, mono_titles_all, "মনপুদি", "poem-mp-extra-1"),
    ("ওলি গীত ( ঘুমপাড়ানি গান)", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-1"),
    ("বিঝু বেড়ান্", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-2"),
    ("কামানি রুয়ে", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-3"),
    ("আয় তুঙ্গোবী", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-4"),
    ("কলিয়ে ভরমেল' পিত্থিমী", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-5"),
    ("অফিস পাড়া", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-6"),
    ("নাদঙ্ ছাড়া", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-7"),
    ("উং চুয়া প্যুঅঃ", tinn_all, tinn_titles_all, "তিন্নোমুরি", "poem-tn-extra-8"),
]

added = 0
for title, all_paras, all_titles, source, poem_id in remaining_poems:
    content_lines = extract_section(all_paras, title, [t for t in all_titles if t != title])
    content = process_content(content_lines)
    excerpt = content[0] if content else title
    
    new_poem = {
        "id": poem_id,
        "title": title,
        "category": "poem",
        "badge": "কবিতা",
        "excerpt": excerpt[:100],
        "date": "তারিখ অজানা",
        "readTime": f"{max(1, len(content)//15)} মিনিট পাঠ",
        "isFeatured": False,
        "content": content,
        "source": source
    }
    data.append(new_poem)
    added += 1
    print(f"  Added: {title} ({len(content)} lines)")

with open('writings.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, separators=(',', ':'))

poems_after = [x for x in data if x.get('category') == 'poem']
print(f"\n✅ Added {added} more poems")
print(f"Total poems on website now: {len(poems_after)}")
print(f"Total entries: {len(data)}")
