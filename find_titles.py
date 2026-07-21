import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

# Find the exact text for these 9 titles in the DOCX files
tinn_doc = docx.Document("তিন্নোমুরি (TINNOMURI).docx")
mono_doc = docx.Document("মনপুদি.docx")

search_titles = [
    "হোচপাঙ্ ভিলিনে", "ওলি গীত", "বিঝু বেড়ান্", "কামানি রুয়ে",
    "আয় তুঙ্গোবী", "কলিয়ে ভরমেল", "অফিস পাড়া", "নাদঙ্", "উং চুয়া"
]

print("=== Searching in তিন্নোমুরি ===")
for para in tinn_doc.paragraphs:
    t = para.text.strip()
    for s in search_titles:
        if s in t:
            print(f"  FOUND: [{t}] (repr: {repr(t[:30])})")
            break

print("\n=== Searching in মনপুদি ===")
for para in mono_doc.paragraphs:
    if para.style.name.startswith('toc'):
        continue
    t = para.text.strip()
    for s in search_titles:
        if s in t:
            print(f"  FOUND: [{t}] (repr: {repr(t[:30])})")
            break
