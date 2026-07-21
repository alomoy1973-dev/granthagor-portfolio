import docx
import json
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

with open('writings.json', encoding='utf-8') as f:
    data = json.load(f)

def process_raw(raw_text):
    """Process raw paragraph text (may have embedded newlines) into clean lines"""
    lines = []
    for line in raw_text.split('\n'):
        line = line.strip()
        if line:
            lines.append(line)
        else:
            if lines and lines[-1] != '__STANZA__':
                lines.append('__STANZA__')
    while lines and lines[-1] == '__STANZA__':
        lines.pop()
    return lines

# =============================================
# Fix হোচপাঙ্ ভিলিনে - too much content
# Remove it and re-add with correct content
# =============================================
data = [x for x in data if not (x.get('title') == 'হোচপাঙ্ ভিলিনে')]

# Read মনপুদি properly
mono_doc = docx.Document("মনপুদি.docx")

# Get all non-TOC paragraphs with their full raw text
mono_raw = [(p.text, p.style.name) for p in mono_doc.paragraphs if not p.style.name.startswith('toc')]

# Identify all TOC titles so we can use them as section boundaries
toc_titles = []
for p in mono_doc.paragraphs:
    if p.style.name.startswith('toc'):
        t = p.text.strip()
        if '\t' in t:
            t = t.split('\t')[0].strip()
        else:
            t = re.sub(r'[\.\s\d]+$', '', t).strip()
        if t and t not in ['মনপুদি', 'আলোময় চাকমা']:
            toc_titles.append(t)

# Build poems map: title -> raw content paragraphs
# In মনপুদি, each title paragraph may contain the title + first line of content
poems_map = {}
current_title = None
current_raw = []

for raw_text, style in mono_raw:
    # Check if this paragraph starts with or equals a known title
    matched = None
    for title in toc_titles:
        if raw_text.strip() == title or raw_text.strip().startswith(title + '\n') or raw_text.strip().startswith(title + ' \n'):
            matched = title
            break
    
    if matched:
        if current_title:
            poems_map[current_title] = current_raw[:]
        current_title = matched
        current_raw = []
        # Get any content after the title in the same paragraph
        remainder = raw_text.strip()
        if remainder.startswith(matched):
            remainder = remainder[len(matched):].strip()
        if remainder:
            current_raw.append(remainder)
    elif current_title and raw_text.strip():
        current_raw.append(raw_text.strip() if raw_text.strip() else '')
    elif current_title:
        current_raw.append('')  # blank line = stanza break

if current_title:
    poems_map[current_title] = current_raw

# Fix poems that have too little content
print("Fixing incomplete মনপুদি poems...")
fixed_count = 0
for i, entry in enumerate(data):
    if entry.get('source') == 'মনপুদি' and len(entry.get('content', [])) <= 2:
        title = entry['title']
        if title in poems_map:
            raw_lines = poems_map[title]
            # Process raw lines (each may contain embedded \n)
            new_content = []
            for raw in raw_lines:
                for line in raw.split('\n'):
                    line = line.strip()
                    if line:
                        new_content.append(line)
                    else:
                        if new_content and new_content[-1] != '__STANZA__':
                            new_content.append('__STANZA__')
            while new_content and new_content[-1] == '__STANZA__':
                new_content.pop()
            
            if len(new_content) > len(entry.get('content', [])):
                data[i]['content'] = new_content
                data[i]['excerpt'] = new_content[0][:100] if new_content else title
                print(f"  Fixed: {title} ({len(entry.get('content',[]))} -> {len(new_content)} lines)")
                fixed_count += 1

# Add হোচপাঙ্ ভিলিনে with correct content
if 'হোচপাঙ্ ভিলিনে' in poems_map:
    raw_lines = poems_map['হোচপাঙ্ ভিলিনে']
    new_content = []
    for raw in raw_lines:
        for line in raw.split('\n'):
            line = line.strip()
            if line:
                new_content.append(line)
            else:
                if new_content and new_content[-1] != '__STANZA__':
                    new_content.append('__STANZA__')
    while new_content and new_content[-1] == '__STANZA__':
        new_content.pop()
    
    new_poem = {
        "id": "poem-mp-n01",
        "title": "হোচপাঙ্ ভিলিনে",
        "category": "poem",
        "badge": "কবিতা",
        "excerpt": new_content[0][:100] if new_content else "হোচপাঙ্ ভিলিনে",
        "date": "তারিখ অজানা",
        "readTime": f"{max(1, len(new_content)//15)} মিনিট পাঠ",
        "isFeatured": False,
        "content": new_content,
        "source": "মনপুদি"
    }
    data.append(new_poem)
    print(f"  Added: হোচপাঙ্ ভিলিনে ({len(new_content)} lines)")
    fixed_count += 1

print(f"\nFixed/added {fixed_count} poems")

# Save
with open('writings.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, separators=(',', ':'))

poems_total = [x for x in data if x.get('category') == 'poem']
print(f"Total poems now: {len(poems_total)}")
print(f"Total entries: {len(data)}")
